# -*- coding: utf-8 -*-
# dashboard/export.py

import io
import csv
import base64
import importlib
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, Query, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from models import get_db, ExportRecord, UserSession, User


def _import_numpy():
    try:
        import numpy as np
        return np
    except ImportError:
        raise HTTPException(500, "numpy not installed. Run: pip install numpy")

def _import_matplotlib():
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        return plt
    except ImportError:
        raise HTTPException(500, "matplotlib not installed. Run: pip install matplotlib")

def _import_reportlab():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, Image as RLImage, PageBreak,
        )
        from reportlab.lib.enums import TA_CENTER
        return {
            "A4": A4, "colors": colors, "getSampleStyleSheet": getSampleStyleSheet,
            "ParagraphStyle": ParagraphStyle, "cm": cm,
            "SimpleDocTemplate": SimpleDocTemplate, "Paragraph": Paragraph,
            "Spacer": Spacer, "Table": Table, "TableStyle": TableStyle,
            "HRFlowable": HRFlowable, "RLImage": RLImage, "PageBreak": PageBreak,
            "TA_CENTER": TA_CENTER,
        }
    except ImportError:
        raise HTTPException(500, "reportlab not installed. Run: pip install reportlab")

def _import_openpyxl():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.drawing.image import Image as XLImage
        return {
            "Workbook": Workbook, "Font": Font, "PatternFill": PatternFill,
            "Alignment": Alignment, "Border": Border, "Side": Side,
            "get_column_letter": get_column_letter, "XLImage": XLImage,
        }
    except ImportError:
        raise HTTPException(500, "openpyxl not installed. Run: pip install openpyxl")

def _import_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return {
            "Document": Document, "Pt": Pt, "RGBColor": RGBColor,
            "Inches": Inches, "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "qn": qn, "OxmlElement": OxmlElement,
        }
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")


# ─────────────────────────────────────────────────────────────
# ✅ UPDATED: ExportRequest now carries project_id + file_upload_id
# ─────────────────────────────────────────────────────────────
class ExportRequest(BaseModel):
    section:        Optional[str]            = "all"
    data:           Optional[Dict[str, Any]] = None
    sections:       Optional[List[str]]      = None
    token:          Optional[str]            = None
    project_id:     Optional[int]            = None   # ✅ NEW
    file_upload_id: Optional[int]            = None   # ✅ NEW


router = APIRouter(prefix="/export", tags=["Export"])


# ─────────────────────────────────────────────────────────────
# ✅ UPDATED: _save_export_record now stores project_id + file_upload_id
# ─────────────────────────────────────────────────────────────
def _save_export_record(
    db:              Session,
    token:           Optional[str],
    export_type:     str,
    filename:        str,
    file_size_bytes: int,
    section:         str,
    project_id:      Optional[int] = None,    # ✅ NEW
    file_upload_id:  Optional[int] = None,    # ✅ NEW
):
    """
    Write one row to export_records.
    Stores: id, user_id, project_id, file_upload_id, export_type,
            export_filename, file_size_kb, status, exported_at.
    Silently no-ops if token is missing/invalid or DB write fails.
    """
    if not token or db is None:
        return
    try:
        session = db.query(UserSession).filter(
            UserSession.token_hash == token,
            UserSession.is_active  == True,
            UserSession.expires_at  > datetime.now(),
        ).first()
        if not session:
            return

        record = ExportRecord(
            user_id         = session.user_id,
            project_id      = project_id,       # ✅ stored
            file_upload_id  = file_upload_id,   # ✅ stored
            export_type     = export_type,
            export_filename = filename,
            file_size_kb    = round(file_size_bytes / 1024, 2),
            status          = "completed",
            download_count  = 1,
            exported_at     = datetime.now(),
        )
        db.add(record)
        db.commit()
    except Exception as e:
        db.rollback()
        print(f"[export] DB write skipped: {e}")


# ─────────────────────────────────────────────────────────────
# Colour tokens
# ─────────────────────────────────────────────────────────────
BG      = "#0F172A"
SURFACE = "#1E293B"
ORANGE  = "#F97316"
BLUE    = "#3B82F6"
GREEN   = "#10B981"
RED     = "#EF4444"
PURPLE  = "#A855F7"
YELLOW  = "#EAB308"
MUTED   = "#64748B"
TEXT    = "#E2E8F0"
PALETTE = [ORANGE, BLUE, GREEN, RED, PURPLE, YELLOW, "#06B6D4", "#EC4899"]
NOW     = lambda: datetime.now().strftime("%Y-%m-%d %H:%M")

# ─────────────────────────────────────────────────────────────
# Data helpers
# ─────────────────────────────────────────────────────────────
def _meta(data):      return (data or {}).get("metadata",    {})
def _stats(data):     return (data or {}).get("statistics",  {})
def _ts(data):        return (data or {}).get("time_series", [])
def _anomalies(data): return (data or {}).get("anomalies",   {})

def _parse_ts(data):
    rows = _ts(data)
    if not rows:
        return [], {}
    time_key = list(rows[0].keys())[0]
    pos_keys = [k for k in rows[0].keys() if k != time_key]
    times  = [float(r.get(time_key, 0)) for r in rows]
    series = {p: [float(r.get(p, 0) or 0) for r in rows] for p in pos_keys}
    return times, series


# ─────────────────────────────────────────────────────────────
# Lazy numpy / matplotlib helpers
# ─────────────────────────────────────────────────────────────
def _np():
    import numpy as np
    return np

def _plt():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    return plt

def _apply_dark_theme(fig, axes_list):
    fig.patch.set_facecolor(BG)
    for ax in axes_list:
        ax.set_facecolor(SURFACE)
        ax.tick_params(colors=TEXT, labelsize=8)
        ax.xaxis.label.set_color(TEXT)
        ax.yaxis.label.set_color(TEXT)
        ax.title.set_color(TEXT)
        for spine in ax.spines.values():
            spine.set_edgecolor(MUTED)
        ax.grid(color=MUTED, linestyle="--", linewidth=0.4, alpha=0.4)


# ─────────────────────────────────────────────────────────────
# Chart generators
# ─────────────────────────────────────────────────────────────
def chart_time_series(data, max_positions=8) -> io.BytesIO:
    times, series = _parse_ts(data)
    if not times:
        return None
    positions = list(series.keys())[:max_positions]
    fig, ax = _plt().subplots(figsize=(11, 4.5))
    _apply_dark_theme(fig, [ax])
    for i, pos in enumerate(positions):
        ax.plot(times, series[pos], color=PALETTE[i % len(PALETTE)],
                linewidth=1.8, label=pos, alpha=0.92)
    ax.set_xlabel("Time (s)", fontsize=9)
    ax.set_ylabel("Temperature (degC)", fontsize=9)
    ax.set_title("Temperature vs Time - All Positions", fontsize=11, fontweight="bold", color=ORANGE)
    ax.legend(fontsize=7, facecolor=SURFACE, edgecolor=MUTED, labelcolor=TEXT, ncol=min(4, len(positions)))
    fig.tight_layout(pad=1.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=BG)
    _plt().close(fig)
    buf.seek(0)
    return buf


def chart_statistics(data) -> io.BytesIO:
    stats = _stats(data)
    if not stats:
        return None
    positions = list(stats.keys())
    means = [stats[p].get("mean", 0) for p in positions]
    stds  = [stats[p].get("std",  0) for p in positions]
    mins  = [stats[p].get("min",  0) for p in positions]
    maxs  = [stats[p].get("max",  0) for p in positions]
    fig, axes = _plt().subplots(1, 2, figsize=(12, 4))
    _apply_dark_theme(fig, axes)
    x = _np().arange(len(positions))
    axes[0].bar(x, means, yerr=stds, color=ORANGE, alpha=0.85,
                error_kw=dict(ecolor=RED, capsize=4, linewidth=1.2), width=0.6)
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(positions, rotation=45, ha="right", fontsize=7)
    axes[0].set_ylabel("Temperature (degC)", fontsize=9)
    axes[0].set_title("Mean Temperature +/- Std Dev", fontsize=10, fontweight="bold", color=ORANGE)
    for i, pos in enumerate(positions):
        axes[1].plot([i, i], [mins[i], maxs[i]], color=BLUE, linewidth=2.5, alpha=0.7)
        axes[1].scatter(i, means[i], color=ORANGE, zorder=5, s=40)
        axes[1].scatter(i, mins[i],  color=GREEN,  zorder=5, s=25, marker="v")
        axes[1].scatter(i, maxs[i],  color=RED,    zorder=5, s=25, marker="^")
    axes[1].set_xticks(range(len(positions)))
    axes[1].set_xticklabels(positions, rotation=45, ha="right", fontsize=7)
    axes[1].set_ylabel("Temperature (degC)", fontsize=9)
    axes[1].set_title("Temperature Range (Min / Mean / Max)", fontsize=10, fontweight="bold", color=ORANGE)
    from matplotlib.lines import Line2D
    legend_el = [
        Line2D([0],[0], marker="^", color="w", markerfacecolor=RED,    markersize=8, label="Max"),
        Line2D([0],[0], marker="o", color="w", markerfacecolor=ORANGE, markersize=8, label="Mean"),
        Line2D([0],[0], marker="v", color="w", markerfacecolor=GREEN,  markersize=8, label="Min"),
    ]
    axes[1].legend(handles=legend_el, fontsize=7, facecolor=SURFACE, edgecolor=MUTED, labelcolor=TEXT)
    fig.tight_layout(pad=1.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=BG)
    _plt().close(fig)
    buf.seek(0)
    return buf


def chart_heatmap(data) -> io.BytesIO:
    times, series = _parse_ts(data)
    if not times or not series:
        return None
    positions = list(series.keys())
    step   = max(1, len(times) // 80)
    t_down = times[::step]
    matrix = _np().array([[series[p][i] for p in positions] for i in range(0, len(times), step)])
    fig, ax = _plt().subplots(figsize=(12, 5))
    _apply_dark_theme(fig, [ax])
    im = ax.imshow(matrix, aspect="auto", origin="lower",
                   extent=[0, len(positions), t_down[0], t_down[-1]],
                   cmap="inferno", interpolation="nearest")
    cbar = fig.colorbar(im, ax=ax, pad=0.02)
    cbar.set_label("Temperature (degC)", color=TEXT, fontsize=9)
    cbar.ax.yaxis.set_tick_params(color=TEXT)
    _plt().setp(cbar.ax.yaxis.get_ticklabels(), color=TEXT, fontsize=8)
    ax.set_xticks(_np().arange(len(positions)) + 0.5)
    ax.set_xticklabels(positions, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("Time (s)", fontsize=9)
    ax.set_title("Thermal Heatmap - Positions x Time", fontsize=11, fontweight="bold", color=ORANGE)
    fig.tight_layout(pad=1.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=BG)
    _plt().close(fig)
    buf.seek(0)
    return buf


def chart_anomalies(data) -> io.BytesIO:
    anomalies = _anomalies(data)
    times, series = _parse_ts(data)
    if not anomalies:
        return None
    positions  = list(anomalies.keys())
    counts     = [anomalies[p].get("count", len(anomalies[p].get("outliers", []))) for p in positions]
    pcts       = [anomalies[p].get("percentage", 0) or 0 for p in positions]
    fig, axes  = _plt().subplots(1, 2, figsize=(12, 4.5))
    _apply_dark_theme(fig, axes)
    bar_colors = [RED if c > 5 else YELLOW if c > 2 else GREEN for c in counts]
    axes[0].bar(positions, counts, color=bar_colors, alpha=0.85, width=0.6)
    axes[0].set_xticklabels(positions, rotation=45, ha="right", fontsize=7)
    axes[0].set_ylabel("Outlier Count", fontsize=9)
    axes[0].set_title("Anomaly Count per Position", fontsize=10, fontweight="bold", color=ORANGE)
    if times and series:
        for i, pos in enumerate(positions[:4]):
            if pos in series:
                axes[1].plot(times, series[pos], color=PALETTE[i], linewidth=1.2, alpha=0.6, label=pos)
                outliers = anomalies[pos].get("outliers", [])
                if outliers:
                    ox = [float(o.get("time",  0) or 0) for o in outliers]
                    oy = [float(o.get("value", 0) or 0) for o in outliers]
                    axes[1].scatter(ox, oy, color=RED, s=50, zorder=10, marker="x", linewidths=2)
        axes[1].set_xlabel("Time (s)", fontsize=9)
        axes[1].set_ylabel("Temperature (degC)", fontsize=9)
        axes[1].set_title("Outliers on Time Series", fontsize=10, fontweight="bold", color=ORANGE)
        axes[1].legend(fontsize=7, facecolor=SURFACE, edgecolor=MUTED, labelcolor=TEXT)
    else:
        axes[1].bar(positions, pcts, color=ORANGE, alpha=0.8, width=0.6)
        axes[1].set_xticklabels(positions, rotation=45, ha="right", fontsize=7)
        axes[1].set_ylabel("Outlier Percentage (%)", fontsize=9)
        axes[1].set_title("Anomaly Percentage per Position", fontsize=10, fontweight="bold", color=ORANGE)
    fig.tight_layout(pad=1.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=BG)
    _plt().close(fig)
    buf.seek(0)
    return buf


def chart_correlation(data) -> io.BytesIO:
    times, series = _parse_ts(data)
    if len(series) < 2:
        return None
    positions = list(series.keys())
    n        = len(positions)
    matrix   = _np().zeros((n, n))
    vals_arr = [_np().array(series[p]) for p in positions]
    for i in range(n):
        for j in range(n):
            if i == j:
                matrix[i][j] = 1.0
            else:
                if vals_arr[i].std() == 0 or vals_arr[j].std() == 0:
                    matrix[i][j] = 0
                else:
                    matrix[i][j] = float(_np().corrcoef(vals_arr[i], vals_arr[j])[0, 1])
    fig, ax = _plt().subplots(figsize=(max(7, n * 0.7), max(6, n * 0.65)))
    _apply_dark_theme(fig, [ax])
    im = ax.imshow(matrix, cmap="RdYlBu_r", vmin=-1, vmax=1, aspect="auto")
    cbar = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.set_label("Pearson r", color=TEXT, fontsize=9)
    cbar.ax.yaxis.set_tick_params(color=TEXT)
    _plt().setp(cbar.ax.yaxis.get_ticklabels(), color=TEXT, fontsize=8)
    ax.set_xticks(range(n))
    ax.set_yticks(range(n))
    ax.set_xticklabels(positions, rotation=45, ha="right", fontsize=max(6, 9 - n // 4))
    ax.set_yticklabels(positions, fontsize=max(6, 9 - n // 4))
    ax.set_title("Pearson Correlation Matrix", fontsize=11, fontweight="bold", color=ORANGE)
    for i in range(n):
        for j in range(n):
            val = matrix[i][j]
            ax.text(j, i, f"{val:.2f}", ha="center", va="center",
                    fontsize=max(5, 8 - n // 5),
                    color="white" if abs(val) > 0.6 else TEXT)
    fig.tight_layout(pad=1.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=BG)
    _plt().close(fig)
    buf.seek(0)
    return buf


def chart_spatial_snapshot(data) -> io.BytesIO:
    times, series = _parse_ts(data)
    if not times or not series:
        return None
    positions = list(series.keys())
    n         = len(times)
    snapshots = {
        f"t = {times[0]:.1f}s (start)": [series[p][0]      for p in positions],
        f"t = {times[n//2]:.1f}s (mid)":[series[p][n // 2] for p in positions],
        f"t = {times[-1]:.1f}s (end)":  [series[p][-1]     for p in positions],
    }
    fig, ax = _plt().subplots(figsize=(11, 4.5))
    _apply_dark_theme(fig, [ax])
    for (label, temps), col in zip(snapshots.items(), [BLUE, ORANGE, RED]):
        ax.plot(positions, temps, color=col, linewidth=2.2, marker="o", markersize=5, label=label)
    ax.set_xlabel("Sensor Position", fontsize=9)
    ax.set_ylabel("Temperature (degC)", fontsize=9)
    ax.set_title("Spatial Temperature Profile - Start / Mid / End", fontsize=11, fontweight="bold", color=ORANGE)
    ax.legend(fontsize=8, facecolor=SURFACE, edgecolor=MUTED, labelcolor=TEXT)
    _plt().xticks(rotation=45, ha="right", fontsize=7)
    fig.tight_layout(pad=1.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=BG)
    _plt().close(fig)
    buf.seek(0)
    return buf


def _build_all_charts(data: dict) -> dict:
    return {
        "Time Series":        chart_time_series(data),
        "Statistics":         chart_statistics(data),
        "Heatmap":            chart_heatmap(data),
        "Anomalies":          chart_anomalies(data),
        "Correlation Matrix": chart_correlation(data),
        "Spatial Profile":    chart_spatial_snapshot(data),
    }


# ═════════════════════════════════════════════════════════════
# 1. CSV
# ═════════════════════════════════════════════════════════════
@router.post("/csv")
@router.get("/csv")
async def export_csv(
    req:     ExportRequest = None,
    section: str           = Query("all"),
    token:   Optional[str] = Query(None),
    db:      Session       = Depends(get_db),
):
    data    = (req.data    or {}) if req else {}
    sec     = (req.section or section) if req else section
    ts      = _ts(data)
    output  = io.StringIO()
    writer  = csv.writer(output)

    if ts:
        headers = list(ts[0].keys())
        writer.writerow(headers)
        for row in ts:
            writer.writerow([row.get(h, "") for h in headers])
    else:
        writer.writerow(["time", "note"])
        writer.writerow(["0", "No data available"])

    content = output.getvalue()
    fname   = f"ThermoPlot_{sec}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

    # ✅ extract project_id + file_upload_id and pass to record helper
    tok            = token or (req.token          if req else None)
    proj_id        = req.project_id               if req else None
    file_upload_id = req.file_upload_id           if req else None
    _save_export_record(db, tok, "csv", fname, len(content.encode()), sec,
                        project_id=proj_id, file_upload_id=file_upload_id)

    return StreamingResponse(
        iter([content]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )


# ═════════════════════════════════════════════════════════════
# 2. EXCEL
# ═════════════════════════════════════════════════════════════
def _xl_hdr(cell, bg="F97316", fg="FFFFFF"):
    from openpyxl.styles import Font, PatternFill, Alignment
    cell.font      = Font(bold=True, color=fg, size=10, name="Calibri")
    cell.fill      = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

def _xl_border(cell):
    from openpyxl.styles import Border, Side
    s = Side(style="thin", color="D1D5DB")
    cell.border = Border(top=s, bottom=s, left=s, right=s)

def _xl_alt(cell, idx):
    from openpyxl.styles import PatternFill
    if idx % 2 == 0:
        cell.fill = PatternFill("solid", fgColor="F8FAFC")

def _auto_width(ws):
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        w = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(w + 4, 45)


@router.post("/excel")
@router.get("/excel")
async def export_excel(
    req:     ExportRequest = None,
    section: str           = Query("all"),
    token:   Optional[str] = Query(None),
    db:      Session       = Depends(get_db),
):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage

    data  = (req.data    or {}) if req else {}
    sec   = (req.section or section) if req else section
    meta  = _meta(data)
    stats = _stats(data)
    ts    = _ts(data)
    anom  = _anomalies(data)

    wb = Workbook()

    # Sheet 1: Overview
    ws = wb.active
    ws.title = "Overview"
    ws.sheet_view.showGridLines = False
    ws["A1"] = "ThermoPlot Analysis Report"
    ws["A1"].font = Font(bold=True, size=16, color="F97316", name="Calibri")
    ws.merge_cells("A1:D1")
    ws["A2"] = f"Generated: {NOW()}"
    ws["A2"].font = Font(italic=True, color="64748B", size=10)
    ws.merge_cells("A2:D2")
    ws.append([])
    ws.append(["Metric", "Value"])
    for cell in ws[ws.max_row]: _xl_hdr(cell)
    tr = meta.get("time_range", [0, 0]) or [0, 0]
    info_rows = [
        ("Section exported",  sec),
        ("Total Time Points", meta.get("rows", "N/A")),
        ("Sensor Positions",  meta.get("position_columns", "N/A")),
        ("Time Range",        f"{tr[0]:.2f}s - {tr[1]:.2f}s"),
        ("Completeness",      f"{meta.get('completeness', 'N/A')}%"),
        ("Missing Values",    meta.get("missing_values", 0)),
    ]
    for i, (k, v) in enumerate(info_rows):
        ws.append([k, v])
        r = ws.max_row
        for c in [ws.cell(r, 1), ws.cell(r, 2)]:
            _xl_alt(c, i); _xl_border(c)
    _auto_width(ws)

    # Sheet 2: Statistics
    if stats:
        ws2 = wb.create_sheet("Statistics")
        ws2.sheet_view.showGridLines = False
        hdrs = ["Position", "Mean (degC)", "Median (degC)", "Std Dev", "Min (degC)", "Max (degC)", "Range (degC)", "CV (%)"]
        ws2.append(hdrs)
        for cell in ws2[1]: _xl_hdr(cell)
        for i, (pos, s) in enumerate(stats.items()):
            ws2.append([pos,
                        round(s.get("mean",   0), 3),
                        round(s.get("median", 0), 3),
                        round(s.get("std",    0), 3),
                        round(s.get("min",    0), 3),
                        round(s.get("max",    0), 3),
                        round(s.get("range",  0), 3),
                        round(s.get("cv",     0), 2)])
            for cell in ws2[ws2.max_row]: _xl_alt(cell, i); _xl_border(cell)
        _auto_width(ws2)

    # Sheet 3: Time Series
    if ts:
        ws3 = wb.create_sheet("Time Series")
        ws3.sheet_view.showGridLines = False
        hdrs = list(ts[0].keys())
        ws3.append(hdrs)
        for cell in ws3[1]: _xl_hdr(cell)
        for i, row in enumerate(ts):
            ws3.append([row.get(h, "") for h in hdrs])
            for cell in ws3[ws3.max_row]: _xl_alt(cell, i); _xl_border(cell)
        _auto_width(ws3)

    # Sheet 4: Anomalies
    if anom:
        ws4 = wb.create_sheet("Anomalies")
        ws4.sheet_view.showGridLines = False
        hdrs = ["Position", "Outlier Count", "Percentage (%)", "Lower Bound (degC)", "Upper Bound (degC)", "Sample Outliers"]
        ws4.append(hdrs)
        for cell in ws4[1]: _xl_hdr(cell, bg="EF4444")
        for i, (pos, a) in enumerate(anom.items()):
            outliers = a.get("outliers", [])
            sample   = ", ".join([f"{o.get('value','?')}@{o.get('time','?')}s" for o in outliers[:4]])
            ws4.append([pos,
                        a.get("count", len(outliers)),
                        a.get("percentage", "N/A"),
                        a.get("lower_bound", "N/A"),
                        a.get("upper_bound", "N/A"),
                        sample or "None"])
            for cell in ws4[ws4.max_row]: _xl_alt(cell, i); _xl_border(cell)
        _auto_width(ws4)

    # Sheet 5: Charts
    ws5 = wb.create_sheet("Charts")
    ws5.sheet_view.showGridLines = False
    ws5["A1"] = "Generated Charts"
    ws5["A1"].font = Font(bold=True, size=14, color="F97316", name="Calibri")
    charts     = _build_all_charts(data)
    row_offset = 2
    for title, buf in charts.items():
        if buf is None:
            continue
        ws5.cell(row=row_offset, column=1, value=title).font = Font(bold=True, size=11, color="0F172A")
        row_offset += 1
        img_buf        = io.BytesIO(buf.read())
        buf.seek(0)
        xl_img         = XLImage(img_buf)
        xl_img.width   = 900
        xl_img.height  = 360
        ws5.add_image(xl_img, f"A{row_offset}")
        row_offset    += 22

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    fname = f"ThermoPlot_{sec}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    # ✅ pass project_id + file_upload_id
    tok            = token or (req.token          if req else None)
    proj_id        = req.project_id               if req else None
    file_upload_id = req.file_upload_id           if req else None
    _save_export_record(db, tok, "excel", fname, out.getbuffer().nbytes, sec,
                        project_id=proj_id, file_upload_id=file_upload_id)
    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )


# ═════════════════════════════════════════════════════════════
# 3. PDF
# ═════════════════════════════════════════════════════════════
def _pdf_tbl_style(hdr_color=None):
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle
    hc = hdr_color or colors.HexColor("#F97316")
    return TableStyle([
        ("BACKGROUND",     (0, 0), (-1, 0),  hc),
        ("TEXTCOLOR",      (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",       (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",       (0, 0), (-1, 0),  8),
        ("ALIGN",          (0, 0), (-1, 0),  "CENTER"),
        ("FONTNAME",       (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE",       (0, 1), (-1, -1), 7.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ("ALIGN",          (1, 1), (-1, -1), "CENTER"),
        ("ALIGN",          (0, 1), (0, -1),  "LEFT"),
        ("GRID",           (0, 0), (-1, -1), 0.35, colors.HexColor("#E2E8F0")),
        ("TOPPADDING",     (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
        ("LEFTPADDING",    (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 6),
    ])

def _buf_to_rl_image(buf: io.BytesIO, width_cm: float, height_cm: float):
    from reportlab.platypus import Image as RLImage
    from reportlab.lib.units import cm
    buf.seek(0)
    return RLImage(buf, width=width_cm * cm, height=height_cm * cm)


@router.post("/pdf")
@router.get("/pdf")
async def export_pdf(
    req:     ExportRequest = None,
    section: str           = Query("all"),
    token:   Optional[str] = Query(None),
    db:      Session       = Depends(get_db),
):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        HRFlowable, Image as RLImage, PageBreak,
    )
    from reportlab.lib.enums import TA_CENTER

    BRAND_ORANGE_RL = colors.HexColor("#F97316")
    BRAND_DARK_RL   = colors.HexColor("#0F172A")
    RED_RL          = colors.HexColor("#EF4444")

    data  = (req.data    or {}) if req else {}
    sec   = (req.section or section) if req else section
    meta  = _meta(data)
    stats = _stats(data)
    ts    = _ts(data)
    anom  = _anomalies(data)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.8*cm, rightMargin=1.8*cm,
                            topMargin=2*cm,    bottomMargin=2*cm)
    SS    = getSampleStyleSheet()
    story = []

    title_s   = ParagraphStyle("tp_t",   parent=SS["Title"],   textColor=BRAND_ORANGE_RL, fontSize=22, spaceAfter=2, fontName="Helvetica-Bold")
    sub_s     = ParagraphStyle("tp_s",   parent=SS["Normal"],  textColor=colors.HexColor("#64748B"), fontSize=9, spaceAfter=10)
    h1_s      = ParagraphStyle("tp_h1",  parent=SS["Heading1"],textColor=BRAND_DARK_RL, fontSize=13, spaceBefore=16, spaceAfter=6, fontName="Helvetica-Bold")
    body_s    = ParagraphStyle("tp_b",   parent=SS["Normal"],  fontSize=8.5, leading=13, textColor=colors.HexColor("#374151"))
    caption_s = ParagraphStyle("tp_cap", parent=SS["Normal"],  fontSize=7.5, leading=11, textColor=colors.HexColor("#64748B"), alignment=TA_CENTER, spaceAfter=8)
    footer_s  = ParagraphStyle("tp_f",   parent=SS["Normal"],  fontSize=7.5, alignment=TA_CENTER, textColor=colors.HexColor("#94A3B8"))

    story.append(Paragraph("ThermoPlot", title_s))
    story.append(Paragraph("Integrated Thermal Research - Analysis Report", sub_s))
    story.append(Paragraph(f"Generated: {NOW()}  .  Section: <b>{sec}</b>", body_s))
    story.append(HRFlowable(width="100%", thickness=2, color=BRAND_ORANGE_RL, spaceAfter=14))

    if meta:
        story.append(Paragraph("Dataset Overview", h1_s))
        tr   = meta.get("time_range", [0, 0]) or [0, 0]
        rows = [
            ["Metric", "Value"],
            ["Total Time Points",  str(meta.get("rows", "N/A"))],
            ["Sensor Positions",   str(meta.get("position_columns", "N/A"))],
            ["Time Range",         f"{tr[0]:.2f}s - {tr[1]:.2f}s"],
            ["Completeness",       f"{meta.get('completeness', 'N/A')}%"],
            ["Missing Values",     str(meta.get("missing_values", 0))],
        ]
        t = Table(rows, colWidths=[8*cm, 9*cm])
        t.setStyle(_pdf_tbl_style())
        story.append(t)
        story.append(Spacer(1, 10))

    ts_chart = chart_time_series(data)
    if ts_chart:
        story.append(Paragraph("Temperature vs Time", h1_s))
        story.append(_buf_to_rl_image(ts_chart, 17, 7))
        story.append(Paragraph("Line chart showing temperature evolution over time for all sensor positions.", caption_s))

    stat_chart = chart_statistics(data)
    if stat_chart:
        story.append(Paragraph("Statistical Summary", h1_s))
        if stats:
            rows = [["Position", "Mean", "Median", "Std Dev", "Min", "Max", "CV%"]]
            for pos, s in stats.items():
                rows.append([pos,
                             f"{s.get('mean',   0):.2f}degC",
                             f"{s.get('median', 0):.2f}degC",
                             f"{s.get('std',    0):.3f}",
                             f"{s.get('min',    0):.2f}degC",
                             f"{s.get('max',    0):.2f}degC",
                             f"{s.get('cv',     0):.1f}%"])
            t = Table(rows, colWidths=[4*cm, 2.2*cm, 2.2*cm, 2.2*cm, 2.2*cm, 2.2*cm, 1.8*cm])
            t.setStyle(_pdf_tbl_style())
            story.append(t)
            story.append(Spacer(1, 8))
        story.append(_buf_to_rl_image(stat_chart, 17, 6.5))
        story.append(Paragraph("Left: Mean +/- Std Dev. Right: Temperature range (min/mean/max) per position.", caption_s))

    hm_chart = chart_heatmap(data)
    if hm_chart:
        story.append(PageBreak())
        story.append(Paragraph("Thermal Heatmap", h1_s))
        story.append(_buf_to_rl_image(hm_chart, 17, 7))
        story.append(Paragraph("Colour-mapped temperature across all sensor positions over time (inferno scale).", caption_s))

    sp_chart = chart_spatial_snapshot(data)
    if sp_chart:
        story.append(Paragraph("Spatial Temperature Profile", h1_s))
        story.append(_buf_to_rl_image(sp_chart, 17, 6.5))
        story.append(Paragraph("Spatial temperature distribution at start, midpoint, and end of the recording.", caption_s))

    an_chart = chart_anomalies(data)
    if anom or an_chart:
        story.append(PageBreak())
        story.append(Paragraph("Anomaly Detection", h1_s))
        story.append(Paragraph("Outliers detected using IQR method: values outside Q1 - 1.5xIQR and Q3 + 1.5xIQR are flagged.", body_s))
        story.append(Spacer(1, 6))
        if anom:
            rows = [["Position", "Count", "Percentage", "Lower Bound", "Upper Bound"]]
            for pos, a in anom.items():
                rows.append([pos,
                             str(a.get("count", len(a.get("outliers", [])))),
                             f"{a.get('percentage', 'N/A')}%",
                             f"{a.get('lower_bound', 'N/A')}degC",
                             f"{a.get('upper_bound', 'N/A')}degC"])
            t = Table(rows, colWidths=[4.5*cm, 2.5*cm, 2.8*cm, 3.2*cm, 3.2*cm])
            t.setStyle(_pdf_tbl_style(hdr_color=RED_RL))
            story.append(t)
            story.append(Spacer(1, 8))
        if an_chart:
            story.append(_buf_to_rl_image(an_chart, 17, 6.5))
            story.append(Paragraph("Left: Outlier count per position. Right: Outlier locations on time series.", caption_s))

    corr_chart = chart_correlation(data)
    if corr_chart:
        story.append(Paragraph("Correlation Matrix", h1_s))
        story.append(_buf_to_rl_image(corr_chart, 17, 7.5))
        story.append(Paragraph("Pearson r correlation between all sensor positions. Red = strong positive, Blue = strong negative.", caption_s))

    story.append(Spacer(1, 18))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0"), spaceAfter=6))
    story.append(Paragraph("Generated by ThermoPlot . Integrated Thermal Research Platform", footer_s))

    doc.build(story)
    buf.seek(0)
    fname = f"ThermoPlot_{sec}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    # ✅ pass project_id + file_upload_id
    tok            = token or (req.token          if req else None)
    proj_id        = req.project_id               if req else None
    file_upload_id = req.file_upload_id           if req else None
    _save_export_record(db, tok, "pdf", fname, buf.getbuffer().nbytes, sec,
                        project_id=proj_id, file_upload_id=file_upload_id)
    buf.seek(0)

    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )


# ═════════════════════════════════════════════════════════════
# 4. WORD
# ═════════════════════════════════════════════════════════════
def _word_shd(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _word_table(doc, headers, rows, hdr_hex="F97316"):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for cell, text in zip(hdr_cells, headers):
        _word_shd(cell, hdr_hex)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold      = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row_data in enumerate(rows):
        cells = tbl.add_row().cells
        bg    = "EFF6FF" if i % 2 == 0 else "FFFFFF"
        for cell, text in zip(cells, row_data):
            _word_shd(cell, bg)
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            run.font.size = Pt(8.5)
    return tbl

def _word_add_chart(doc, buf: io.BytesIO, caption: str, width_inches=6.2):
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if buf is None:
        return
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size      = Pt(8)
    cap.runs[0].font.italic    = True
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()


@router.post("/word")
@router.get("/word")
async def export_word(
    req:     ExportRequest = None,
    section: str           = Query("all"),
    token:   Optional[str] = Query(None),
    db:      Session       = Depends(get_db),
):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data  = (req.data    or {}) if req else {}
    sec   = (req.section or section) if req else section
    meta  = _meta(data)
    stats = _stats(data)
    ts    = _ts(data)
    anom  = _anomalies(data)

    doc      = Document()
    doc_sec  = doc.sections[0]
    doc_sec.page_width    = Inches(8.5)
    doc_sec.page_height   = Inches(11)
    doc_sec.left_margin   = doc_sec.right_margin  = Inches(1)
    doc_sec.top_margin    = doc_sec.bottom_margin = Inches(1)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr_run = tp.add_run("ThermoPlot - Analysis Report")
    tr_run.font.size = Pt(22); tr_run.font.bold = True
    tr_run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"Generated: {NOW()}  .  Section: {sec}")
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()

    if meta:
        h = doc.add_heading("Dataset Overview", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        tr_range = meta.get("time_range", [0, 0]) or [0, 0]
        _word_table(doc, ["Metric", "Value"], [
            ["Total Time Points",  str(meta.get("rows", "N/A"))],
            ["Sensor Positions",   str(meta.get("position_columns", "N/A"))],
            ["Time Range",         f"{tr_range[0]:.2f}s - {tr_range[1]:.2f}s"],
            ["Completeness",       f"{meta.get('completeness', 'N/A')}%"],
            ["Missing Values",     str(meta.get("missing_values", 0))],
        ])
        doc.add_paragraph()

    ts_chart = chart_time_series(data)
    if ts_chart:
        h = doc.add_heading("Temperature vs Time", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        _word_add_chart(doc, ts_chart, "Line chart: temperature over time for all sensor positions.")

    stat_chart = chart_statistics(data)
    if stats or stat_chart:
        h = doc.add_heading("Statistical Summary", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        if stats:
            _word_table(doc,
                ["Position", "Mean (degC)", "Median", "Std Dev", "Min", "Max", "CV%"],
                [[pos,
                  f"{s.get('mean',   0):.2f}",
                  f"{s.get('median', 0):.2f}",
                  f"{s.get('std',    0):.3f}",
                  f"{s.get('min',    0):.2f}",
                  f"{s.get('max',    0):.2f}",
                  f"{s.get('cv',     0):.1f}%"]
                 for pos, s in stats.items()])
            doc.add_paragraph()
        if stat_chart:
            _word_add_chart(doc, stat_chart, "Left: Mean +/- Std Dev. Right: Min / Mean / Max range per position.")

    hm_chart = chart_heatmap(data)
    if hm_chart:
        h = doc.add_heading("Thermal Heatmap", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        _word_add_chart(doc, hm_chart, "Colour-mapped temperature across all positions over time (inferno scale).")

    sp_chart = chart_spatial_snapshot(data)
    if sp_chart:
        h = doc.add_heading("Spatial Temperature Profile", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        _word_add_chart(doc, sp_chart, "Spatial temperature distribution at start, midpoint, and end of the recording.")

    an_chart = chart_anomalies(data)
    if anom or an_chart:
        h = doc.add_heading("Anomaly Detection", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        note = doc.add_paragraph()
        nr   = note.add_run("IQR method: values outside Q1 - 1.5xIQR  to  Q3 + 1.5xIQR are flagged.")
        nr.font.size = Pt(9); nr.font.italic = True
        nr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        if anom:
            _word_table(doc,
                ["Position", "Count", "Percentage", "Lower Bound", "Upper Bound"],
                [[pos,
                  str(a.get("count", len(a.get("outliers", [])))),
                  f"{a.get('percentage', 'N/A')}%",
                  f"{a.get('lower_bound', 'N/A')}degC",
                  f"{a.get('upper_bound', 'N/A')}degC"]
                 for pos, a in anom.items()],
                hdr_hex="EF4444")
            doc.add_paragraph()
        if an_chart:
            _word_add_chart(doc, an_chart, "Left: Outlier count per position. Right: Outlier locations on time series.")

    corr_chart = chart_correlation(data)
    if corr_chart:
        h = doc.add_heading("Correlation Matrix", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        _word_add_chart(doc, corr_chart, "Pearson r between all sensor positions. Red = positive, Blue = negative correlation.")

    doc.add_paragraph()
    fp = doc.add_paragraph("Generated by ThermoPlot . Integrated Thermal Research Platform")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    fname = f"ThermoPlot_{sec}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    # ✅ pass project_id + file_upload_id
    tok            = token or (req.token          if req else None)
    proj_id        = req.project_id               if req else None
    file_upload_id = req.file_upload_id           if req else None
    _save_export_record(db, tok, "word", fname, out.getbuffer().nbytes, sec,
                        project_id=proj_id, file_upload_id=file_upload_id)
    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )